import tkinter as tk
from tkinter import messagebox, ttk, simpledialog
from PIL import Image, ImageTk
import sqlite3
import datetime
import docx

class DoiBong:
    def __init__(self, MaDB, TenDoi, NamTL, GiaTriDB):
        self.MaDB = MaDB
        self.TenDoi = TenDoi
        self.NamTL = NamTL
        self.GiaTriDB = GiaTriDB

class CauThu:
    def __init__(self, MaCT, TenCauThu, MaDB, TenDoi, SoAo, ViTri, NgaySinh, QuocTich, ChieuCao, CanNang, GiaTriCT):
        self.MaCT = MaCT
        self.TenCauThu = TenCauThu
        self.MaDB = MaDB
        self.TenDoi = TenDoi
        self.SoAo = SoAo
        self.ViTri = ViTri
        self.NgaySinh = NgaySinh
        self.QuocTich = QuocTich
        self.ChieuCao = ChieuCao
        self.CanNang = CanNang
        self.GiaTriCT = GiaTriCT

class ChuyenDich:
    def __init__(self, MaGD, MaCT, MaDB, LichSu):
        self.MaGD = MaGD
        self.MaCT = MaCT
        self.MaDB = MaDB
        self.LichSu = LichSu

class HienThiCauThu:
    def __init__(self, ma_db):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()
        self.ma_db = ma_db

    def hien_thi(self):
        self.cursor.execute("SELECT * FROM CauThu WHERE MaDB = ?", (self.ma_db,))
        cau_thu = self.cursor.fetchall()

        if cau_thu:
            print("Danh sách cầu thủ:")
            for ct in cau_thu:
                print(f"""
Mã CT: {ct[0]}
Tên: {ct[1]}
Mã Đội: {ct[2]}
Tên Đội: {ct[3]}
Số áo: {ct[4]}
Vị trí: {ct[5]}
Ngày sinh: {ct[6]}
Quốc tịch: {ct[7]}
Chiều cao: {ct[8]} cm
Cân nặng: {ct[9]} kg
Giá trị: {ct[10]}
""")
        else:
            print("Không có cầu thủ nào trong cơ sở dữ liệu.")

class ThemCauThu:
    def __init__(self, ma_db, ten_db):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()
        self.ma_db = ma_db
        self.ten_db = ten_db

    def nhap_thong_tin(self):
        while True:
            # Không cần nhập MaCT nữa
            self.TenCauThu = input("Nhập tên cầu thủ (hoặc 'q' để thoát): ")
            if self.TenCauThu == 'q':
                break
            self.SoAo = int(input("Nhập số áo: "))
            self.ViTri = input("Nhập vị trí: ")
            self.NgaySinh = input("Nhập ngày sinh (YYYY-MM-DD): ")
            self.QuocTich = input("Nhập quốc tịch: ")
            self.ChieuCao = float(input("Nhập chiều cao (cm): "))
            self.CanNang = float(input("Nhập cân nặng (kg): "))
            self.GiaTriCT = float(input("Nhập giá trị cầu thủ: "))

            self.them_cau_thu()

    def them_cau_thu(self):
        try:
            # Không cần đặt giá trị cho MaCT nữa
            self.cursor.execute("""
                INSERT INTO CauThu (TenCauThu, MaDB, TenDoi, SoAo, ViTri, NgaySinh, QuocTich, ChieuCao, CanNang, GiaTriCT)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (self.TenCauThu, self.ma_db, self.ten_db, self.SoAo, self.ViTri, self.NgaySinh, self.QuocTich, self.ChieuCao, self.CanNang, self.GiaTriCT))
            self.conn.commit()
            print("Thêm cầu thủ thành công!")
        except sqlite3.Error as e:
            print("Lỗi khi thêm cầu thủ:", e)

class XoaCauThu:
    def __init__(self, ma_db):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()
        self.ma_db = ma_db

    def xoa_cau_thu(self, ma_ct):  # Nhận ma_ct làm tham số
        try:
            self.cursor.execute("DELETE FROM CauThu WHERE MaCT = ?", (ma_ct,))  # Xóa cầu thủ dựa trên ma_ct
            self.conn.commit()
            if self.cursor.rowcount > 0:
                print("Xóa cầu thủ thành công!")
            else:
                print("Không tìm thấy cầu thủ với mã CT đó.")
        except sqlite3.Error as e:
            print("Lỗi khi xóa cầu thủ:", e)

class SuaCauThu:
    def __init__(self, ma_db):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()
        self.ma_db = ma_db

    def sua_cau_thu(self):
        self.hien_thi = HienThiCauThu(self.ma_db)
        self.hien_thi.hien_thi()
        ma_ct = input("Nhập mã cầu thủ để sửa (hoặc 'q' để thoát): ")
        if ma_ct == 'q':
            return

        self.cursor.execute("SELECT * FROM CauThu WHERE MaCT = ? AND MaDB = ?", (ma_ct, self.ma_db))
        cau_thu = self.cursor.fetchone()

        if cau_thu:
            # Nhận thông tin mới từ người dùng
            ten_moi = input(f"Nhập tên mới (hiện tại: {cau_thu[1]}): ")
            so_ao_moi = int(input(f"Nhập số áo mới (hiện tại: {cau_thu[4]}): "))
            vi_tri_moi = input(f"Nhập vị trí mới (hiện tại: {cau_thu[5]}): ")
            ngay_sinh_moi = input(f"Nhập ngày sinh mới (YYYY-MM-DD, hiện tại: {cau_thu[6]}): ")
            quoc_tich_moi = input(f"Nhập quốc tịch mới (hiện tại: {cau_thu[7]}): ")
            chieu_cao_moi = float(input(f"Nhập chiều cao mới (cm, hiện tại: {cau_thu[8]}): "))
            can_nang_moi = float(input(f"Nhập cân nặng mới (kg, hiện tại: {cau_thu[9]}): "))
            gia_tri_moi = float(input(f"Nhập giá trị mới (hiện tại: {cau_thu[10]}): "))

            try:
                self.cursor.execute("""
                    UPDATE CauThu
                    SET TenCauThu = ?, SoAo = ?, ViTri = ?, NgaySinh = ?, QuocTich = ?, ChieuCao = ?, CanNang = ?, GiaTriCT = ?
                    WHERE MaCT = ? AND MaDB = ?
                """, (ten_moi, so_ao_moi, vi_tri_moi, ngay_sinh_moi, quoc_tich_moi, chieu_cao_moi, can_nang_moi, gia_tri_moi, ma_ct, self.ma_db))
                self.conn.commit()
                print("Cập nhật thông tin cầu thủ thành công!")
            except sqlite3.Error as e:
                print("Lỗi khi sửa thông tin cầu thủ:", e)
        else:
            print("Không tìm thấy cầu thủ với mã CT đó.")

    def __del__(self):
        self.conn.close()

class HienThiDoiBong:
    def __init__(self):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()

    def hien_thi(self):
        self.cursor.execute("SELECT * FROM DoiBong")
        doi_bong = self.cursor.fetchall()

        if doi_bong:
            print("Danh sách đội bóng:")
            for db in doi_bong:
                print(f"""
Mã Đội: {db[0]}
Tên Đội: {db[1]}
Năm Thành Lập: {db[2]}
""")
        else:
            print("Không có đội bóng nào trong cơ sở dữ liệu.")

class ThemDoiBong:
    def __init__(self):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()

    def them_doi_bong(self):
        try:
            self.cursor.execute("""
                INSERT INTO DoiBong (TenDoi, NamTL)
                VALUES (?, ?)
            """, (self.TenDoi, self.NamTL))
            self.conn.commit()
            print("Thêm đội bóng thành công!")
        except sqlite3.Error as e:
            print("Lỗi khi thêm đội bóng:", e)

class XoaDoiBong:
    def __init__(self):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()

    def xoa_doi_bong(self, ma_db):
        # Xóa cầu thủ thuộc đội bóng
        self.cursor.execute("DELETE FROM CauThu WHERE MaDB = ?", (ma_db,))
        # Xóa đội bóng
        self.cursor.execute("DELETE FROM DoiBong WHERE MaDB = ?", (ma_db,))
        self.conn.commit()

class SuaDoiBong:
    def __init__(self):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()

    def sua_doi_bong(self, ma_db, ten_doi_moi, nam_tl_moi):  # Thêm 2 tham số ten_doi_moi và nam_tl_moi
        try:
            # Cập nhật thông tin đội bóng
            self.cursor.execute("""
                UPDATE DoiBong
                SET TenDoi = ?, NamTL = ?
                WHERE MaDB = ?
            """, (ten_doi_moi, nam_tl_moi, ma_db))
            self.conn.commit()

            # Cập nhật thông tin đội bóng của cầu thủ
            self.cursor.execute("""
                UPDATE CauThu
                SET TenDoi = ?
                WHERE MaDB = ?
            """, (ten_doi_moi, ma_db))
            self.conn.commit()
            print("Cập nhật thông tin đội bóng thành công!")

        except sqlite3.Error as e:
            print("Lỗi khi sửa thông tin đội bóng:", e)

def tinh_tong_gia_tri_doi_bong(ma_db):
    conn = sqlite3.connect('bongda.db')
    cursor = conn.cursor()
    cursor.execute("SELECT SUM(GiaTriCT) FROM CauThu WHERE MaDB = ?", (ma_db,))
    tong_gia_tri = cursor.fetchone()[0]
    conn.close()
    return tong_gia_tri if tong_gia_tri is not None else 0  # Tránh trường hợp None

class TimKiemCauThu:
    def __init__(self, ma_db, hien_thi_callback):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()  # Thêm dòng này
        self.ma_db = ma_db
        self.hien_thi_callback = hien_thi_callback

    def tim_kiem(self):
        # ... (Thay thế toàn bộ phần code trong phương thức tim_kiem bằng code sau)
        search_window = tk.Toplevel()
        search_window.title("Tìm Kiếm Cầu Thủ")

        search_label = tk.Label(search_window, text="Tìm kiếm theo:")
        search_label.grid(row=0, column=0, padx=5, pady=5)

        search_var = tk.IntVar(value=1)  # Mặc định chọn tìm theo tên
        search_options = [
            ("Tên Cầu Thủ", 1),
            ("Vị Trí", 2),
            ("Quốc Tịch", 3)
        ]
        for text, val in search_options:
            tk.Radiobutton(search_window, text=text, variable=search_var, value=val).grid(row=1+val, column=0, sticky="w")

        keyword_label = tk.Label(search_window, text="Nhập từ khóa:")
        keyword_label.grid(row=5, column=0, padx=5, pady=5)
        entry_keyword = tk.Entry(search_window)
        entry_keyword.grid(row=6, column=0, padx=5, pady=5)

        def thuc_hien_tim_kiem():
            lua_chon = search_var.get()
            keyword = entry_keyword.get()

            if lua_chon == 1:
                sql = "SELECT * FROM CauThu WHERE TenCauThu LIKE ? AND MaDB = ?"
                self.cursor.execute(sql, ('%' + keyword + '%', self.ma_db))
            elif lua_chon == 2:
                sql = "SELECT * FROM CauThu WHERE ViTri LIKE ? AND MaDB = ?"
                self.cursor.execute(sql, ('%' + keyword + '%', self.ma_db))
            elif lua_chon == 3:
                sql = "SELECT * FROM CauThu WHERE QuocTich LIKE ? AND MaDB = ?"
                self.cursor.execute(sql, ('%' + keyword + '%', self.ma_db))

            ket_qua = self.cursor.fetchall()
            if ket_qua:
                self.hien_thi_callback(ket_qua)  # Truyền kết quả vào callback nếu có kết quả
            else:
                messagebox.showinfo("Thông báo", "Không tìm thấy cầu thủ nào phù hợp.") # Thông báo nếu không tìm thấy
            search_window.destroy()

        search_button = tk.Button(search_window, text="Tìm Kiếm", command=thuc_hien_tim_kiem)
        search_button.grid(row=7, column=0, columnspan=2, pady=10)

class SapXepCauThu:
    def __init__(self, ma_db, hien_thi_callback):
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()
        self.ma_db = ma_db
        self.hien_thi_callback = hien_thi_callback
        self.sort_var = tk.IntVar(value=1)  # Thêm dòng này
    
    def sap_xep(self):
        sort_window = tk.Toplevel()
        sort_window.title("Sắp Xếp Cầu Thủ")

        sort_label = tk.Label(sort_window, text="Sắp xếp theo giá trị:")
        sort_label.grid(row=0, column=0, padx=5, pady=5)

        sort_options = [
            ("Tăng dần", 1),
            ("Giảm dần", 2)
        ]
        for text, val in sort_options:
            tk.Radiobutton(sort_window, text=text, variable=self.sort_var, value=val).grid(row=1+val, column=0, sticky="w")  

        sort_button = tk.Button(sort_window, text="Sắp Xếp", command=lambda: self.thuc_hien_sap_xep(sort_window)) # Sửa dòng này
        sort_button.grid(row=4, column=0, columnspan=2, pady=10)

    def thuc_hien_sap_xep(self, sort_window):
        lua_chon = self.sort_var.get()

        if lua_chon == 1:
            sql = "SELECT * FROM CauThu WHERE MaDB = ? ORDER BY GiaTriCT ASC"
            self.cursor.execute(sql, (self.ma_db,))  # Di chuyển dòng này vào trong khối if
            ket_qua = self.cursor.fetchall()
            self.hien_thi_callback(ket_qua)
            sort_window.destroy()
        elif lua_chon == 2:
            sql = "SELECT * FROM CauThu WHERE MaDB = ? ORDER BY GiaTriCT DESC"
            self.cursor.execute(sql, (self.ma_db,))  # Di chuyển dòng này vào trong khối elif
            ket_qua = self.cursor.fetchall()
            self.hien_thi_callback(ket_qua)
            sort_window.destroy()

import sqlite3
from tkinter import messagebox
import datetime

def chuyen_nhuong_cau_thu(ma_ct, ma_db_moi, doi_bong_moi, ten_cau_thu, player_window):
    conn = sqlite3.connect('bongda.db')
    cursor = conn.cursor()

    try:
        # Lấy thông tin đội bóng cũ
        cursor.execute("SELECT TenDoi FROM CauThu WHERE MaCT = ?", (ma_ct,))
        ten_db_cu = cursor.fetchone()[0]  # Lấy tên đội bóng cũ

        # Cập nhật thông tin cầu thủ
        cursor.execute("""
            UPDATE CauThu
            SET MaDB = ?, TenDoi = ?
            WHERE MaCT = ?
        """, (ma_db_moi, doi_bong_moi, ma_ct))
        conn.commit()

        # Lấy thông tin đội bóng cũ (bổ sung)
        cursor.execute("SELECT MaDB FROM CauThu WHERE MaCT = ?", (ma_ct,))
        ma_db_cu = cursor.fetchone()[0]  # Lấy MaDB của đội bóng cũ

        # Lưu lịch sử chuyển nhượng (thay đổi MaDB thành MaDBCu)
        ngay_thang_nam = datetime.datetime.now().strftime("%d/%m/%Y")
        cursor.execute("""
            INSERT INTO ChuyenDich (MaCT, MaDB, TenCauThu, GiaTriCT, TenDoi, NgayGiaoDich)
            VALUES (?, ?, ?, ?, ?, ?)
        """, (ma_ct, ma_db_cu, ten_cau_thu, player_window.player_entries[9].get(), ten_db_cu, ngay_thang_nam))
        conn.commit()

        # Hiển thị thông báo
        messagebox.showinfo("Thông báo", f"Chuyển nhượng {ten_cau_thu} từ {ten_db_cu} sang {doi_bong_moi} thành công!")

        # Cập nhật Listbox
        player_window.hien_thi_cau_thu()

    except sqlite3.Error as e:
        messagebox.showerror("Lỗi", f"Lỗi khi chuyển nhượng: {e}")

    conn.close()


class PlayerManagementWindow(tk.Toplevel):
    # ... (Các phương thức khác)

    def ban_cau_thu(self):
        selection = self.player_listbox.curselection()
        if selection:
            index = selection[0]
            ma_ct = self.player_listbox.get(index).split(" - ")[0]  # Lấy Mã CT
            self.cursor.execute("SELECT TenCauThu FROM CauThu WHERE MaDB = ? AND MaCT = ?", (self.ma_db, ma_ct))
            ten_cau_thu = self.cursor.fetchone()[0]

            # Tạo cửa sổ chuyển nhượng
            transfer_window = tk.Toplevel(self)
            transfer_window.transient(self)
            transfer_window.title("Chuyển nhượng cầu thủ")
            transfer_window.configure(bg="#CCEEFF")

            # Lấy danh sách đội bóng khác
            self.cursor.execute("SELECT TenDoi FROM DoiBong WHERE MaDB != ?", (self.ma_db,))
            doi_bong_khac = [row[0] for row in self.cursor.fetchall()]

            # Tạo OptionMenu chọn đội bóng
            selected_team = tk.StringVar(transfer_window)
            team_option = ttk.OptionMenu(transfer_window, selected_team, *doi_bong_khac)
            team_option.pack(pady=10)

            def chuyen_nhuong():
                doi_bong_moi = selected_team.get()
                if doi_bong_moi:
                    chuyen_nhuong_cau_thu(ma_ct, None, doi_bong_moi, ten_cau_thu, self) # Gọi hàm chuyển nhượng
                    transfer_window.destroy()
                else:
                    messagebox.showwarning("Lỗi", "Vui lòng chọn đội bóng chuyển nhượng.")

            # Nút "Chuyển nhượng"
            transfer_button = tk.Button(transfer_window, text="Chuyển nhượng", command=chuyen_nhuong,
                                      bg="#99FFCC", fg="white", font=("Arial", 12))
            transfer_button.pack(pady=5)
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn cầu thủ để chuyển nhượng.")

def xem_lich_su_chuyen_nhuong(ma_db):
    conn = sqlite3.connect('bongda.db')
    cursor = conn.cursor()
    cursor.execute("""
        SELECT * FROM ChuyenDich 
        WHERE MaDB = ? 
    """, (ma_db,)) # Chỉ cần MaDB thôi
    lich_su = cursor.fetchall()
    conn.close()

    history_window = tk.Toplevel()
    history_window.title("Lịch Sử Chuyển Nhượng")

    if lich_su:
        # Tạo Treeview để hiển thị lịch sử (bỏ MaDBCu)
        tree = ttk.Treeview(history_window, columns=("MaGD", "MaCT", "MaDB", "TenCauThu", "GiaTriCT", "TenDoi", "NgayGiaoDich"), show="headings")
        tree.heading("MaGD", text="Mã Giao Dịch")
        tree.heading("MaCT", text="Mã Cầu Thủ")
        tree.heading("MaDB", text="Mã Đội Bóng Cũ")  # Thay đổi tiêu đề cột
        tree.heading("TenCauThu", text="Tên Cầu Thủ")
        tree.heading("GiaTriCT", text="Giá Trị")
        tree.heading("TenDoi", text="Tên Đội Bóng Cũ")
        tree.heading("NgayGiaoDich", text="Ngày Giao Dịch")

        # Thêm dữ liệu vào Treeview
        for gd in lich_su:
            tree.insert("", "end", values=gd)

        tree.pack(expand=True, fill="both")

        # Frame chứa nút "In lịch sử" và "Xóa lịch sử"
        button_frame = tk.Frame(history_window)
        button_frame.pack(pady=10)

        # Nút "In lịch sử"
        def in_lich_su():
            doc = docx.Document()
            doc.add_heading("Lịch Sử Chuyển Nhượng", 0)

            # Tạo bảng trong Word
            table = doc.add_table(rows=1, cols=len(tree["columns"]))
            table.style = 'Table Grid'  # Thêm style cho bảng

            # Định dạng tiêu đề cột
            hdr_cells = table.rows[0].cells
            for i, col in enumerate(tree["columns"]):
                hdr_cells[i].text = tree.heading(col, "text")
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True  # In đậm tiêu đề cột

            # Thêm dữ liệu vào bảng
            for row in tree.get_children():
                row_data = tree.item(row, "values")
                row_cells = table.add_row().cells
                for i, data in enumerate(row_data):
                    row_cells[i].text = str(data)

            # Lưu file Word
            doc.save("lich_su_chuyen_nhuong.docx")
            messagebox.showinfo("Thông báo", "In lịch sử chuyển nhượng thành công!")

        print_button = tk.Button(button_frame, text="In lịch sử", command=in_lich_su)
        print_button.pack(side="left", padx=5)

        # Nút "Xóa lịch sử"
        def xoa_lich_su():
            if messagebox.askyesno("Xác nhận xóa", "Bạn có chắc chắn muốn xóa lịch sử chuyển nhượng?"):
                xoa_lich_su_chuyen_nhuong(ma_db)
                history_window.destroy()

        delete_button = tk.Button(button_frame, text="Xóa lịch sử", command=xoa_lich_su)
        delete_button.pack(side="left", padx=5)

    else:
        tk.Label(history_window, text="Không có lịch sử chuyển nhượng.").pack(pady=20)

def xoa_lich_su_chuyen_nhuong(ma_db):
    conn = sqlite3.connect('bongda.db')
    cursor = conn.cursor()
    cursor.execute("""
        DELETE FROM ChuyenDich 
        WHERE MaDB = ?
    """, (ma_db,)) # Chỉ cần MaDB thôi
    conn.commit()
    print("Xóa lịch sử chuyển nhượng thành công!")
    conn.close()

class LoginApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Login")
        self.geometry('600x400')
        self.configure(bg='#CCEEFF')  # Màu xanh lá pastel
        self.create_widgets()

    def create_widgets(self):
        image = Image.open("bongda.jpg")
        photo = ImageTk.PhotoImage(image)

        canvas = tk.Canvas(self, width=600, height=400)
        canvas.pack(fill="none", expand=True)
        canvas.create_image(0, 0, image=photo, anchor="nw")

        frame = tk.Frame(self, bg='#CCEEFF')
        canvas.create_window(300, 200, window=frame)

        login_label = tk.Label(frame, text="Đăng Nhập", fg='#3C7363', bg="#CCEEFF", font=('aria', 23))
        username_label = tk.Label(frame, text="Tài Khoản", fg='gray1', bg="#CCEEFF", font=('aria', 13))
        self.username_entry = tk.Entry(frame, fg='gray1', bg='snow', font=('aria', 13))
        password_label = tk.Label(frame, text="Mật Khẩu", fg='gray1', bg="#CCEEFF", font=('aria', 13))
        self.password_entry = tk.Entry(frame, show="*", bg='snow', fg='gray1', font=('aria', 13))
        login_button = tk.Button(frame, text="Đăng Nhập", fg='#3C7363', bg="#CCEEFF", 
                               font=('aria', 13), command=self.login)

        login_label.grid(row=0, column=0, columnspan=2, sticky='news', pady=35)
        username_label.grid(row=1, column=0)
        self.username_entry.grid(row=1, column=1, pady=5)
        password_label.grid(row=2, column=0)
        self.password_entry.grid(row=2, column=1, pady=5)
        login_button.grid(row=3, column=0, columnspan=2, pady=15)

    def login(self):
        username = "admin"
        password = "123"
        if self.username_entry.get() == username and self.password_entry.get() == password:
            messagebox.showinfo(title='Login Success', message="You successfully logged in.")
            self.destroy()
            # Hiển thị giao diện Team Management sau khi đăng nhập thành công
            show_team_management()
        else:
            messagebox.showerror(title="Error", message="Invalid login.")

class TeamManagementApp(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Quản Lý Đội Bóng")
        self.geometry("400x300")
        self.configure(bg="#CCEEFF")  # Màu xanh lá pastel
        self.conn = sqlite3.connect('bongda.db') # Kết nối database
        self.cursor = self.conn.cursor()
        self.create_widgets()

    def create_widgets(self):
        # Tạo frame chứa các nút
        button_frame = tk.Frame(self, bg="#CCEEFF")
        button_frame.pack(pady=20)

        # Nút "Tạo đội bóng"
        create_button = tk.Button(button_frame, text="Tạo đội bóng", command=self.create_team,
                                   bg="#99FFCC", fg="white", font=("Arial", 12))
        create_button.pack(pady=5)

        # Nút "Sửa đội bóng"
        edit_button = tk.Button(button_frame, text="Sửa đội bóng", command=self.edit_team,
                                 bg="#99FFCC", fg="white", font=("Arial", 12))
        edit_button.pack(pady=5)

        # Nút "Xóa đội bóng"
        delete_button = tk.Button(button_frame, text="Xóa đội bóng", command=self.delete_team,
                                   bg="#99FFCC", fg="white", font=("Arial", 12))
        delete_button.pack(pady=5)

        # Frame chứa OptionMenu và nút "Vào đội bóng"
        self.join_frame = tk.Frame(self, bg="#CCEEFF")
        self.join_frame.pack()

        # OptionMenu hiển thị danh sách đội bóng
        self.selected_team = tk.StringVar(self)
        self.team_option = ttk.OptionMenu(self.join_frame, self.selected_team, "") 
        self.team_option.config(width=20)
        self.team_option.grid(row=0, column=0, padx=5)

        # Nút "Vào đội bóng"
        join_button = tk.Button(self.join_frame, text="Vào đội bóng", command=self.join_team,
                                bg="#99FFCC", fg="white", font=("Arial", 12))
        join_button.grid(row=0, column=1, padx=5)

        # Lấy danh sách đội bóng từ database và cập nhật OptionMenu
        self.refresh_team_list()

    def refresh_team_list(self):
        """Cập nhật danh sách đội bóng trong OptionMenu."""
        self.cursor.execute("SELECT TenDoi FROM DoiBong")
        teams = [row[0] for row in self.cursor.fetchall()]
        self.team_option['menu'].delete(0, 'end') 
        for team in teams:
            self.team_option['menu'].add_command(label=team, command=tk._setit(self.selected_team, team))
        if teams:  # Chọn đội đầu tiên nếu có
            self.selected_team.set(teams[0])

    def create_team(self):
        # Sử dụng simpledialog để lấy thông tin đội bóng
        ten_doi = simpledialog.askstring("Tạo đội bóng", "Nhập tên đội:")
        if ten_doi is None:
            return  # Người dùng nhấn Cancel

        # Tạo cửa sổ Toplevel tạm thời
        temp_window = tk.Toplevel(self)
        temp_window.withdraw()  # Ẩn cửa sổ tạm thời

        nam_tl = simpledialog.askinteger("Tạo đội bóng", "Nhập năm thành lập:", parent=temp_window)  # Sử dụng temp_window làm parent
        if nam_tl is None:
            temp_window.destroy()  # Hủy cửa sổ tạm thời
            return  # Người dùng nhấn Cancel

        temp_window.destroy()  # Hủy cửa sổ tạm thời

        try:
            them_doi_bong = ThemDoiBong()
            them_doi_bong.TenDoi = ten_doi
            them_doi_bong.NamTL = nam_tl
            them_doi_bong.them_doi_bong()
            self.refresh_team_list()  # Cập nhật lại OptionMenu
            messagebox.showinfo("Thông báo", "Tạo đội bóng thành công!")
        except sqlite3.Error as e:
            messagebox.showerror("Lỗi", f"Lỗi khi tạo đội bóng: {e}")
    
    def edit_team(self):
        try:
            selected_team = self.selected_team.get()
            if not selected_team:
                messagebox.showwarning("Lỗi", "Vui lòng chọn đội bóng để sửa.")
                return

            # Lấy mã đội bóng từ tên đội bóng đã chọn
            self.cursor.execute("SELECT MaDB FROM DoiBong WHERE TenDoi = ?", (selected_team,))
            ma_db = self.cursor.fetchone()
            if not ma_db:
                messagebox.showerror("Lỗi", f"Không tìm thấy đội bóng '{selected_team}' trong cơ sở dữ liệu.")
                return
            ma_db = ma_db[0]

            # Hiển thị hộp thoại để người dùng nhập thông tin mới
            ten_doi_moi = tk.simpledialog.askstring("Sửa đội bóng", "Nhập tên đội bóng mới:", initialvalue=selected_team)
            if not ten_doi_moi:
                return  # Người dùng nhấn Cancel

            nam_tl_moi = tk.simpledialog.askinteger("Sửa đội bóng", "Nhập năm thành lập mới:")
            if nam_tl_moi is None:
                return  # Người dùng nhấn Cancel

            # Cập nhật thông tin đội bóng trong database
            sua_doi_bong = SuaDoiBong()
            sua_doi_bong.sua_doi_bong(ma_db, ten_doi_moi, nam_tl_moi)

            # Cập nhật lại OptionMenu
            self.refresh_team_list()

            messagebox.showinfo("Thông báo", "Sửa đội bóng thành công!")
        except sqlite3.Error as e:
            messagebox.showerror("Lỗi", f"Lỗi khi sửa đội bóng: {e}")

    def delete_team(self):
        try:
            selected_team = self.selected_team.get()
            if not selected_team:
                messagebox.showwarning("Lỗi", "Vui lòng chọn đội bóng để xóa.")
                return

            # Xác nhận xóa
            confirm = messagebox.askyesno("Xác nhận xóa", f"Bạn có chắc chắn muốn xóa đội bóng '{selected_team}'?\n"
                                        "Thao tác này sẽ xóa tất cả cầu thủ thuộc đội bóng này.")
            if not confirm:
                return

            # Lấy mã đội bóng từ tên đội bóng đã chọn
            self.cursor.execute("SELECT MaDB FROM DoiBong WHERE TenDoi = ?", (selected_team,))
            ma_db = self.cursor.fetchone()
            if not ma_db:
                messagebox.showerror("Lỗi", f"Không tìm thấy đội bóng '{selected_team}' trong cơ sở dữ liệu.")
                return
            ma_db = ma_db[0]

            # Xóa đội bóng
            xoa_doi_bong = XoaDoiBong()
            xoa_doi_bong.xoa_doi_bong(ma_db)

            # Cập nhật lại OptionMenu
            self.refresh_team_list()

            messagebox.showinfo("Thông báo", "Xóa đội bóng thành công!")
        except sqlite3.Error as e:
            messagebox.showerror("Lỗi", f"Lỗi khi xóa đội bóng: {e}")

    def join_team(self):
        selected_team = self.selected_team.get()
        if selected_team:
            try:
                # Lấy mã đội bóng từ tên đội bóng đã chọn
                self.cursor.execute("SELECT MaDB FROM DoiBong WHERE TenDoi = ?", (selected_team,))
                ma_db = self.cursor.fetchone()
                if ma_db:
                    ma_db = ma_db[0]
                    # Tạo cửa sổ quản lý cầu thủ
                    PlayerManagementWindow(self, ma_db, selected_team).mainloop() 
                else:
                    messagebox.showwarning("Lỗi", f"Không tìm thấy đội bóng '{selected_team}' trong cơ sở dữ liệu.")
            except sqlite3.Error as e:
                messagebox.showerror("Lỗi", f"Lỗi khi lấy mã đội bóng: {e}")

class PlayerManagementWindow(tk.Toplevel):
    def __init__(self, parent, ma_db, ten_doi):
        super().__init__(parent)
        self.title(f"Quản Lý Cầu Thủ - {ten_doi}")
        self.ma_db = ma_db
        self.configure(bg="#CCEEFF")

        # Khởi tạo thuộc tính cursor
        self.conn = sqlite3.connect('bongda.db')
        self.cursor = self.conn.cursor()

        # --- Frame tiêu đề ---
        title_frame = tk.Frame(self, bg="#CCEEFF")
        title_frame.pack(fill="x", pady=10)

        # Tên đội bóng ở giữa
        team_name_label = tk.Label(title_frame, text=ten_doi, font=("Arial", 18, "bold"), bg="#FFCC99", relief="solid", borderwidth=2)
        team_name_label.pack(side="left", expand=True)

        # Nút tìm kiếm
        search_button = tk.Button(title_frame, text="Tìm kiếm", bg="#99FFCC", fg="white", font=("Arial", 12), command=self.tim_kiem_cau_thu)
        search_button.pack(side="right", padx=5)

        # Nút sắp xếp
        sort_button = tk.Button(title_frame, text="Sắp xếp", bg="#99FFCC", fg="white", font=("Arial", 12), command=self.sap_xep_cau_thu)
        sort_button.pack(side="right", padx=5)

        # --- Frame giá trị đội hình ---
        value_frame = tk.Frame(self, bg="#CCEEFF")
        value_frame.pack(fill="x", pady=5)

        self.total_value_label = tk.Label(value_frame, text="Giá trị đội hình hiện tại: 0 $", bg="#CCEEFF")
        self.total_value_label.pack()

        # --- Frame chứa Listbox và thông tin cầu thủ ---
        content_frame = tk.Frame(self, bg="#CCEEFF")
        content_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # --- Listbox danh sách cầu thủ ---
        player_list_frame = tk.Frame(content_frame, bg="#CCEEFF")
        player_list_frame.pack(side="left", fill="both", expand=True)

        self.player_listbox = tk.Listbox(player_list_frame, width=30)
        self.player_listbox.pack(fill="both", expand=True)
        self.player_listbox.bind("<<ListboxSelect>>", self.on_player_select)

        # --- Frame thông tin cầu thủ ---
        player_info_frame = tk.LabelFrame(content_frame, text="Thông tin cầu thủ", bg="#CCEEFF")
        player_info_frame.pack(side="right", fill="both", expand=True, padx=10)

        labels = ["Mã CT:", "Tên:", "Tên Đội:", "Số áo:", "Vị trí:", "Ngày sinh:", "Quốc tịch:", "Chiều cao:", "Cân nặng:", "Giá trị:"]
        self.player_entries = []  # Danh sách các Entry để hiển thị thông tin cầu thủ

        for i, label_text in enumerate(labels):
            label = tk.Label(player_info_frame, text=label_text, bg="#CCEEFF")
            label.grid(row=i, column=0, padx=5, pady=2, sticky="w")

            entry = tk.Entry(player_info_frame, state="readonly")  # Entry chỉ đọc
            entry.grid(row=i, column=1, padx=5, pady=2)
            self.player_entries.append(entry)

        # --- Frame chức năng ---
        function_frame = tk.Frame(self, bg="#CCEEFF")
        function_frame.pack(pady=10)

        buttons = ["Lịch sử chuyển nhượng", "In danh sách\ncầu thủ", "Thống kê", "Thêm", "Xóa", "Sửa", "Bán"]
        self.function_buttons = []  # Danh sách các nút chức năng

        for btn_text in buttons:
            button = tk.Button(function_frame, text=btn_text, bg="#99FFCC", fg="white", font=("Arial", 12))
            button.pack(side="left", padx=5)
            self.function_buttons.append(button)

        # Vô hiệu hóa các nút "Xóa", "Sửa", "Bán" ban đầu
        self.function_buttons[4]['state'] = tk.DISABLED
        self.function_buttons[5]['state'] = tk.DISABLED
        self.function_buttons[6]['state'] = tk.DISABLED

        # --- Kết nối các nút chức năng với backend ---
        self.function_buttons[0].config(command=self.xem_lich_su_chuyen_nhuong)
        self.function_buttons[1].config(command=self.in_danh_sach_cau_thu)
        self.function_buttons[2].config(command=self.thong_ke)
        self.function_buttons[3].config(command=self.them_cau_thu)
        self.function_buttons[4].config(command=self.xoa_cau_thu)
        self.function_buttons[5].config(command=self.sua_cau_thu)
        self.function_buttons[6].config(command=self.ban_cau_thu)

        # --- Hiển thị danh sách cầu thủ ---
        self.hien_thi_cau_thu()

    def hien_thi_cau_thu(self, cau_thu=None):  # Thêm tham số cau_thu
        self.player_listbox.delete(0, tk.END)
        self.clear_player_info()

        if cau_thu is None:  # Nếu không có tham số cau_thu, lấy tất cả cầu thủ
            self.cursor.execute("SELECT * FROM CauThu WHERE MaDB = ?", (self.ma_db,))
            cau_thu = self.cursor.fetchall()

        for player in cau_thu:
            self.player_listbox.insert(tk.END, f"{player[0]} - {player[1]}")  # Hiển thị Mã CT và tên cầu thủ
        self.cap_nhat_tong_gia_tri()
        self.player_listbox.bind("<<ListboxSelect>>", self.on_player_select)
        
    def cap_nhat_tong_gia_tri(self):
        tong_gia_tri = tinh_tong_gia_tri_doi_bong(self.ma_db)
        self.total_value_label.config(text=f"Giá trị đội hình hiện tại: {tong_gia_tri} $")

    def on_player_select(self, event):
        selection = self.player_listbox.curselection()
        if selection:
            index = selection[0]
            ma_ct = self.player_listbox.get(index).split(" - ")[0]  # Lấy Mã CT từ Listbox
            
            # Lấy thông tin cầu thủ từ database dựa trên Mã CT
            self.cursor.execute("SELECT * FROM CauThu WHERE MaDB = ? AND MaCT = ?", (self.ma_db, ma_ct))
            selected_player = self.cursor.fetchone()

            # Hiển thị thông tin cầu thủ lên các Entry
            self.player_entries[0].config(state="normal")
            self.player_entries[0].delete(0, tk.END)
            self.player_entries[0].insert(0, selected_player[0])  # Mã CT
            self.player_entries[0].config(state="readonly")

            self.player_entries[1].config(state="normal")
            self.player_entries[1].delete(0, tk.END)
            self.player_entries[1].insert(0, selected_player[1])  # Tên
            self.player_entries[1].config(state="readonly")

            self.player_entries[2].config(state="normal")
            self.player_entries[2].delete(0, tk.END)
            self.player_entries[2].insert(0, selected_player[3])  # Tên Đội
            self.player_entries[2].config(state="readonly")

            self.player_entries[3].config(state="normal")
            self.player_entries[3].delete(0, tk.END)
            self.player_entries[3].insert(0, selected_player[4])  # Số áo
            self.player_entries[3].config(state="readonly")

            self.player_entries[4].config(state="normal")
            self.player_entries[4].delete(0, tk.END)
            self.player_entries[4].insert(0, selected_player[5])  # Vị trí
            self.player_entries[4].config(state="readonly")

            self.player_entries[5].config(state="normal")
            self.player_entries[5].delete(0, tk.END)
            self.player_entries[5].insert(0, selected_player[6])  # Ngày sinh
            self.player_entries[5].config(state="readonly")

            self.player_entries[6].config(state="normal")
            self.player_entries[6].delete(0, tk.END)
            self.player_entries[6].insert(0, selected_player[7])  # Quốc tịch
            self.player_entries[6].config(state="readonly")

            self.player_entries[7].config(state="normal")
            self.player_entries[7].delete(0, tk.END)
            self.player_entries[7].insert(0, selected_player[8])  # Chiều cao
            self.player_entries[7].config(state="readonly")

            self.player_entries[8].config(state="normal")
            self.player_entries[8].delete(0, tk.END)
            self.player_entries[8].insert(0, selected_player[9])  # Cân nặng
            self.player_entries[8].config(state="readonly")

            self.player_entries[9].config(state="normal")
            self.player_entries[9].delete(0, tk.END)
            self.player_entries[9].insert(0, selected_player[10])  # Giá trị
            self.player_entries[9].config(state="readonly")

            # Kích hoạt các nút "Xóa", "Sửa", "Bán"
            self.function_buttons[4]['state'] = tk.NORMAL
            self.function_buttons[5]['state'] = tk.NORMAL
            self.function_buttons[6]['state'] = tk.NORMAL
            
    def tim_kiem_cau_thu(self):
        tim_kiem_cau_thu = TimKiemCauThu(self.ma_db, self.hien_thi_cau_thu)
        tim_kiem_cau_thu.tim_kiem()

    def sap_xep_cau_thu(self):
        sap_xep_cau_thu = SapXepCauThu(self.ma_db, self.hien_thi_cau_thu)
        sap_xep_cau_thu.sap_xep()

    # --- Các hàm xử lý sự kiện cho nút ---
    def xem_lich_su_chuyen_nhuong(self):
        xem_lich_su_chuyen_nhuong(self.ma_db)

    def in_danh_sach_cau_thu(self):
        # TODO: Thêm logic in danh sách cầu thủ
        print("In danh sách cầu thủ")

    def thong_ke(self):
        # TODO: Thêm logic thống kê
        print("Thống kê")

    def them_cau_thu(self):
        # Tạo cửa sổ Toplevel để nhập thông tin cầu thủ
        add_player_window = tk.Toplevel(self)
        add_player_window.transient(self)  # Giữ cửa sổ con luôn nằm trước cửa sổ cha
        add_player_window.title("Thêm Cầu Thủ")
        add_player_window.configure(bg="#CCEEFF")

        # Tạo các label và entry cho thông tin cầu thủ
        labels = ["Tên:", "Số áo:", "Vị trí:", "Ngày sinh (YYYY-MM-DD):", "Quốc tịch:", "Chiều cao (cm):", "Cân nặng (kg):", "Giá trị:"]
        entries = []
        for i, label_text in enumerate(labels):
            label = tk.Label(add_player_window, text=label_text, bg="#CCEEFF")
            label.grid(row=i, column=0, padx=5, pady=5, sticky="w")
            entry = tk.Entry(add_player_window)
            entry.grid(row=i, column=1, padx=5, pady=5)
            entries.append(entry)

        def them_cau_thu_vao_db():
            # Lấy thông tin từ các entry
            ten_cau_thu = entries[0].get()
            so_ao = entries[1].get()
            vi_tri = entries[2].get()
            ngay_sinh = entries[3].get()
            quoc_tich = entries[4].get()
            chieu_cao = entries[5].get()
            can_nang = entries[6].get()
            gia_tri = entries[7].get()

            # Kiểm tra xem người dùng đã nhập đủ thông tin chưa
            if not all([ten_cau_thu, so_ao, vi_tri, ngay_sinh, quoc_tich, chieu_cao, can_nang, gia_tri]):
                messagebox.showwarning("Thiếu thông tin", "Vui lòng nhập đầy đủ thông tin.")
                return

            try:
                so_ao = int(so_ao)
                chieu_cao = float(chieu_cao)
                can_nang = float(can_nang)
                gia_tri = float(gia_tri)
            except ValueError:
                messagebox.showwarning("Lỗi", "Số áo, chiều cao, cân nặng và giá trị phải là số.")
                return

            try:
                them_cau_thu = ThemCauThu(self.ma_db, self.title().split(" - ")[1])
                them_cau_thu.TenCauThu = ten_cau_thu
                them_cau_thu.SoAo = so_ao
                them_cau_thu.ViTri = vi_tri
                them_cau_thu.NgaySinh = ngay_sinh
                them_cau_thu.QuocTich = quoc_tich
                them_cau_thu.ChieuCao = chieu_cao
                them_cau_thu.CanNang = can_nang
                them_cau_thu.GiaTriCT = gia_tri
                them_cau_thu.them_cau_thu()
                self.hien_thi_cau_thu()  # Cập nhật Listbox sau khi thêm
                add_player_window.destroy()  # Đóng cửa sổ sau khi thêm
                messagebox.showinfo("Thông báo", "Thêm cầu thủ thành công!")
            except sqlite3.Error as e:
                messagebox.showerror("Lỗi", f"Lỗi khi thêm cầu thủ: {e}")

        # Nút "Thêm"
        add_button = tk.Button(add_player_window, text="Thêm", command=them_cau_thu_vao_db,
                              bg="#99FFCC", fg="white", font=("Arial", 12))
        add_button.grid(row=len(labels), column=0, columnspan=2, padx=5, pady=10)
    
    def sua_cau_thu(self):
        selection = self.player_listbox.curselection()
        if selection:
            index = selection[0]
            self.cursor.execute("SELECT * FROM CauThu WHERE MaDB = ?", (self.ma_db,))
            player_data = self.cursor.fetchall()[index]
            ma_ct = player_data[0]

            # Tạo cửa sổ Toplevel để sửa thông tin cầu thủ
            edit_player_window = tk.Toplevel(self)
            edit_player_window.transient(self)
            edit_player_window.title("Sửa Thông Tin Cầu Thủ")
            edit_player_window.configure(bg="#CCEEFF")

            # Tạo các label và entry cho thông tin cầu thủ
            labels = ["Tên:", "Số áo:", "Vị trí:", "Ngày sinh (YYYY-MM-DD):", "Quốc tịch:", "Chiều cao (cm):", "Cân nặng (kg):", "Giá trị:"]
            entries = []
            for i, label_text in enumerate(labels):
                label = tk.Label(edit_player_window, text=label_text, bg="#CCEEFF")
                label.grid(row=i, column=0, padx=5, pady=5, sticky="w")
                entry = tk.Entry(edit_player_window)
                entry.grid(row=i, column=1, padx=5, pady=5)
                # Bỏ phần tự động điền thông tin, để người dùng nhập
                # entry.insert(0, player_data[i + 1])
                entries.append(entry)

            def luu_thong_tin_cau_thu():
                # Lấy thông tin từ các entry
                ten_cau_thu = entries[0].get()
                so_ao = entries[1].get()
                vi_tri = entries[2].get()
                ngay_sinh = entries[3].get()
                quoc_tich = entries[4].get()
                chieu_cao = entries[5].get()
                can_nang = entries[6].get()
                gia_tri = entries[7].get()

                # Kiểm tra xem người dùng đã nhập đủ thông tin chưa
                if not all([ten_cau_thu, so_ao, vi_tri, ngay_sinh, quoc_tich, chieu_cao, can_nang, gia_tri]):
                    messagebox.showwarning("Thiếu thông tin", "Vui lòng nhập đầy đủ thông tin.")
                    return

                try:
                    so_ao = int(so_ao)
                    chieu_cao = float(chieu_cao)
                    can_nang = float(can_nang)
                    gia_tri = float(gia_tri)
                except ValueError:
                    messagebox.showwarning("Lỗi", "Số áo, chiều cao, cân nặng và giá trị phải là số.")
                    return

                try:
                    # Cập nhật thông tin cầu thủ trong database
                    self.cursor.execute("""
                        UPDATE CauThu
                        SET TenCauThu = ?, SoAo = ?, ViTri = ?, NgaySinh = ?, QuocTich = ?, ChieuCao = ?, CanNang = ?, GiaTriCT = ?
                        WHERE MaCT = ?
                    """, (ten_cau_thu, so_ao, vi_tri, ngay_sinh, quoc_tich, chieu_cao, can_nang, gia_tri, ma_ct))
                    self.conn.commit()
                    self.hien_thi_cau_thu()
                    edit_player_window.destroy()
                    messagebox.showinfo("Thông báo", "Sửa thông tin cầu thủ thành công!")
                except sqlite3.Error as e:
                    messagebox.showerror("Lỗi", f"Lỗi khi sửa thông tin cầu thủ: {e}")

            # Nút "Lưu"
            save_button = tk.Button(edit_player_window, text="Lưu", command=luu_thong_tin_cau_thu,
                                  bg="#99FFCC", fg="white", font=("Arial", 12))
            save_button.grid(row=len(labels), column=0, columnspan=2, padx=5, pady=10)
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn cầu thủ để sửa.")

    def xoa_cau_thu(self):
        selection = self.player_listbox.curselection()
        if selection:
            index = selection[0]
            ma_ct = self.player_listbox.get(index).split(" - ")[0]  # Lấy Mã CT từ Listbox

            # Lấy tên cầu thủ từ database dựa trên Mã CT đã lấy
            self.cursor.execute("SELECT TenCauThu FROM CauThu WHERE MaDB = ? AND MaCT = ?", (self.ma_db, ma_ct))
            ten_cau_thu = self.cursor.fetchone()[0]

            # Xác nhận xóa
            confirm = messagebox.askyesno("Xác nhận xóa", f"Bạn có chắc chắn muốn xóa cầu thủ '{ten_cau_thu}'?")
            if confirm:
                xoa_cau_thu = XoaCauThu(self.ma_db)
                xoa_cau_thu.xoa_cau_thu(ma_ct)  # Truyền ma_ct vào hàm xoa_cau_thu
                self.hien_thi_cau_thu()
                self.clear_player_info()
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn cầu thủ để xóa.")

    def ban_cau_thu(self):
        selection = self.player_listbox.curselection()
        if selection:
            index = selection[0]
            self.cursor.execute("SELECT MaCT, TenCauThu FROM CauThu WHERE MaDB = ?", (self.ma_db,))
            ma_ct, ten_cau_thu = self.cursor.fetchall()[index]

            # Tạo cửa sổ Toplevel để chọn đội bóng chuyển nhượng
            transfer_window = tk.Toplevel(self)
            transfer_window.transient(self)
            transfer_window.title("Chuyển nhượng cầu thủ")
            transfer_window.configure(bg="#CCEEFF")

            # Lấy danh sách đội bóng từ database
            self.cursor.execute("SELECT TenDoi FROM DoiBong WHERE MaDB != ?", (self.ma_db,))
            doi_bong_khac = [row[0] for row in self.cursor.fetchall()]

            # Tạo OptionMenu để chọn đội bóng
            selected_team = tk.StringVar(transfer_window)
            team_option = ttk.OptionMenu(transfer_window, selected_team, *doi_bong_khac)
            team_option.pack(pady=10)

            def chuyen_nhuong():
                doi_bong_moi = selected_team.get()
                if doi_bong_moi:
                    try:
                        # Lấy MaDB của đội bóng mới
                        self.cursor.execute("SELECT MaDB FROM DoiBong WHERE TenDoi = ?", (doi_bong_moi,))
                        ma_db_moi = self.cursor.fetchone()[0]

                        # Thực hiện chuyển nhượng
                        chuyen_nhuong_cau_thu(ma_ct, ma_db_moi, doi_bong_moi, ten_cau_thu, self)
                        self.hien_thi_cau_thu()
                        transfer_window.destroy()
                    except sqlite3.Error as e:
                        messagebox.showerror("Lỗi", f"Lỗi khi chuyển nhượng: {e}")
                else:
                    messagebox.showwarning("Lỗi", "Vui lòng chọn đội bóng chuyển nhượng.")

            # Nút "Chuyển nhượng"
            transfer_button = tk.Button(transfer_window, text="Chuyển nhượng", command=chuyen_nhuong,
                                      bg="#99FFCC", fg="white", font=("Arial", 12))
            transfer_button.pack(pady=5)
        else:
            messagebox.showwarning("Lỗi", "Vui lòng chọn cầu thủ để chuyển nhượng.")
            
    def clear_player_info(self):
        # Xóa thông tin cầu thủ trên các Entry
        for entry in self.player_entries:
            entry.config(state="normal")
            entry.delete(0, tk.END)
            entry.config(state="readonly")

    def tim_kiem_cau_thu(self):
        tim_kiem_cau_thu = TimKiemCauThu(self.ma_db, self.hien_thi_cau_thu)  # Thêm self.hien_thi_cau_thu
        tim_kiem_cau_thu.tim_kiem()

    def sap_xep_cau_thu(self):
        sap_xep_cau_thu = SapXepCauThu(self.ma_db, self.hien_thi_cau_thu)  # Thêm self.hien_thi_cau_thu
        sap_xep_cau_thu.sap_xep()

# Hàm để hiển thị giao diện Team Management
def show_team_management():
    TeamManagementApp().mainloop()  # Khởi tạo và chạy vòng lặp chính


# Biến điều kiện để kiểm tra xem có cần hiển thị login hay không
show_login = True

if __name__ == "__main__":
    if show_login:
        login_app = LoginApp()
        login_app.mainloop()
    else:
        show_team_management()
